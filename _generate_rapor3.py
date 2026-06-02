from pathlib import Path
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = Path(r"c:\Users\Beyza\OneDrive - sakarya.edu.tr\Masaüstü\yolarkadasim (2)\yolarkadasim")
md_path = base / "RAPOR_4_FINAL_PROJE_RAPORU.md"
docx_path = base / "RAPOR_3_FORMAT_RAPORU.docx"

text = md_path.read_text(encoding="utf-8")
lines = text.splitlines()

# Create document
doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)
for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
    try:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(size)
    except Exception:
        pass

# Cover page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("YOL ARKADAŞIM\n")
r.bold = True
r.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("AI Destekli Seyahat Planlama Uygulaması\n").italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Final Proje Raporu - Rapor 3 Formatı\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"Tarih: {datetime.now().strftime('%d %B %Y')}\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Sakarya Üniversitesi\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Proje Ekibi: Yol Arkadaşım Geliştiricileri\n")

# Page break
doc.add_page_break()

# Table of contents placeholder (auto-generated in Word by the user)
doc.add_paragraph('İÇİNDEKİLER', style='Heading 1')
# Extract top-level headings from md to populate TOC
heading_re = re.compile(r"^##+\s+(.*)")
headings = []
for line in lines:
    m = heading_re.match(line)
    if m:
        headings.append(m.group(1).strip())

for i, h in enumerate(headings, start=1):
    p = doc.add_paragraph()
    p.style = 'List Number'
    p.add_run(h)

doc.add_page_break()

# Convert markdown to Word with heading mapping
bullet_re = re.compile(r"^[-*]\s+")
num_re = re.compile(r"^\d+[.)]\s+")

current_para = None
for raw in lines:
    line = raw.rstrip()
    stripped = line.strip()
    if not stripped:
        current_para = None
        continue
    if stripped.startswith('---'):
        current_para = None
        continue
    if stripped.startswith('# '):
        doc.add_paragraph(stripped[2:].strip(), style='Title')
        current_para = None
        continue
    if stripped.startswith('## '):
        doc.add_paragraph(stripped[3:].strip(), style='Heading 1')
        current_para = None
        continue
    if stripped.startswith('### '):
        doc.add_paragraph(stripped[4:].strip(), style='Heading 2')
        current_para = None
        continue
    if stripped.startswith('#### '):
        doc.add_paragraph(stripped[5:].strip(), style='Heading 3')
        current_para = None
        continue
    if bullet_re.match(stripped):
        txt = bullet_re.sub('', stripped)
        doc.add_paragraph(txt, style='List Bullet')
        current_para = None
        continue
    if num_re.match(stripped):
        txt = num_re.sub('', stripped)
        doc.add_paragraph(txt, style='List Number')
        current_para = None
        continue
    if stripped.startswith('```'):
        current_para = None
        continue
    # normal paragraph
    if current_para is None:
        current_para = doc.add_paragraph()
        current_para.paragraph_format.space_after = Pt(6)
        current_para.paragraph_format.line_spacing = 1.15
    else:
        current_para.add_run(' ')
    current_para.add_run(stripped)

# Save document
doc.save(docx_path)
print('Saved:', docx_path)
